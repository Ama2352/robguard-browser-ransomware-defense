"""
Generate realistic-looking demo files for the RøB attack demonstration.
Run this script to (re)create the victim folder contents before each demo.

Usage:
    python generate.py

Output: 8 files in the same directory as this script.
"""

import os, openpyxl
from pathlib import Path

HERE = Path(__file__).parent

# ---------------------------------------------------------------------------
# TXT files
# ---------------------------------------------------------------------------

(HERE / "Q1_Financial_Summary.txt").write_text("""\
AFSAM — Q1 2024 Financial Summary
===================================

Revenue
  Product Sales        $1,284,500
  Service Contracts      $342,800
  Licensing Fees         $198,000
  Total Revenue        $1,825,300

Expenses
  Salaries & Benefits    $620,000
  Infrastructure         $145,200
  Marketing               $89,400
  R&D                    $210,000
  Total Expenses       $1,064,600

Net Income               $760,700
Operating Margin           41.7%

Notes:
- Product sales up 12% YoY driven by enterprise segment.
- Infrastructure costs include new cloud migration expenses.
- R&D budget increased by 15% for upcoming platform redesign.
- Q2 forecast: $1.95M revenue, pending contract renewals.

Prepared by: Finance Department
Date: 2024-03-31
""", encoding="utf-8")

(HERE / "Meeting_Notes_March.txt").write_text("""\
Team Meeting Notes — March 15, 2024
=====================================
Attendees: Sarah K., James T., Minh H., Laura P., David R.

Agenda Items
------------

1. Q1 Progress Review
   - Sales target at 94% completion with two weeks remaining.
   - Customer churn reduced from 8.2% to 5.7% this quarter.
   - Action: James to prepare final Q1 report by March 28.

2. Platform Outage Post-Mortem (March 8)
   - Root cause: database connection pool exhaustion under load spike.
   - Fix deployed March 9. SLA breach acknowledged; credits issued.
   - Action: David to implement connection pool monitoring by April 1.

3. Upcoming Product Launch — AFSAM v3.0
   - Beta testing begins April 10 with 20 selected enterprise clients.
   - Key features: AI-assisted anomaly detection, bulk export, SSO.
   - Action: Minh to coordinate beta onboarding materials by April 5.

4. Hiring Updates
   - Two backend engineer offers accepted, start dates in April.
   - Still searching for a senior UX designer.

Next Meeting: April 5, 2024 @ 10:00 AM
""", encoding="utf-8")

# ---------------------------------------------------------------------------
# XLSX files
# ---------------------------------------------------------------------------

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

def make_budget_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "2024 Budget"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2F5496")
    center = Alignment(horizontal="center")

    headers = ["Department", "Q1 Budget", "Q1 Actual", "Q2 Budget", "Q3 Budget", "Q4 Budget", "Annual Total"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center

    rows = [
        ("Engineering",    320000, 308500, 335000, 335000, 340000, 1330000),
        ("Sales",          180000, 192400, 190000, 195000, 210000,  775000),
        ("Marketing",       89000,  87200,  92000,  88000,  95000,  364000),
        ("Operations",      75000,  74100,  76000,  76000,  78000,  305000),
        ("HR & Admin",      42000,  41800,  43000,  43000,  44000,  172000),
        ("R&D",            210000, 215600, 220000, 225000, 230000,  885000),
        ("Customer Success", 55000, 53900,  57000,  57000,  60000,  229000),
    ]

    for row_idx, row in enumerate(rows, 2):
        for col_idx, val in enumerate(row, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = 16

    wb.save(HERE / "Budget_2024.xlsx")

def make_employee_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Employee Records"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="375623")

    headers = ["ID", "Name", "Department", "Role", "Start Date", "Salary", "Status"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill

    employees = [
        ("E001", "Sarah Kim",      "Engineering",    "CTO",                "2018-03-01", 185000, "Active"),
        ("E002", "James Tran",     "Finance",        "CFO",                "2019-07-15", 175000, "Active"),
        ("E003", "Minh Hoang",     "Engineering",    "Lead Backend Eng.",  "2020-01-10", 142000, "Active"),
        ("E004", "Laura Patel",    "Sales",          "Sales Director",     "2019-11-20", 135000, "Active"),
        ("E005", "David Reyes",    "Operations",     "DevOps Engineer",    "2021-04-05", 118000, "Active"),
        ("E006", "Anna Schmidt",   "Marketing",      "Marketing Manager",  "2020-08-17", 112000, "Active"),
        ("E007", "Kevin Okafor",   "Engineering",    "Senior Frontend",    "2021-09-01", 125000, "Active"),
        ("E008", "Priya Sharma",   "Customer Succ.", "CS Lead",            "2022-02-14", 105000, "Active"),
        ("E009", "Tom Nguyen",     "R&D",            "ML Engineer",        "2022-06-01", 138000, "Active"),
        ("E010", "Chen Wei",       "Engineering",    "Backend Engineer",   "2023-01-16", 115000, "Active"),
    ]

    for row_idx, emp in enumerate(employees, 2):
        for col_idx, val in enumerate(emp, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = 20

    wb.save(HERE / "Employee_Records.xlsx")

make_budget_xlsx()
make_employee_xlsx()

# ---------------------------------------------------------------------------
# DOCX files
# ---------------------------------------------------------------------------

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def make_proposal_docx():
    doc = Document()

    title = doc.add_heading("Project Proposal: AFSAM Platform v3.0", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph("Prepared by: Product Team | Date: April 2024 | Version: 1.2")
    doc.add_paragraph("")

    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(
        "This proposal outlines the development roadmap for AFSAM Platform v3.0, "
        "targeting enterprise clients requiring AI-assisted financial statement analysis. "
        "The upgrade focuses on three pillars: intelligent anomaly detection, bulk data "
        "processing, and enhanced security compliance."
    )

    doc.add_heading("2. Objectives", level=2)
    for obj in [
        "Reduce manual review time by 60% through AI-assisted flagging.",
        "Support bulk ingestion of up to 10,000 documents per session.",
        "Achieve SOC 2 Type II compliance by Q3 2024.",
        "Integrate SSO with Okta, Azure AD, and Google Workspace.",
        "Deliver mobile-responsive dashboard for on-the-go reporting.",
    ]:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_heading("3. Timeline", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Phase", "Duration", "Deliverable"
    phases = [
        ("Discovery & Design", "Apr 1 – Apr 30", "UI/UX mockups, architecture docs"),
        ("Core Development",   "May 1 – Jul 31", "AI engine, bulk importer, SSO"),
        ("Beta Testing",       "Aug 1 – Sep 15", "Feedback from 20 enterprise clients"),
        ("Launch",             "Oct 1",           "GA release + documentation"),
    ]
    for phase in phases:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = phase

    doc.add_heading("4. Budget Estimate", level=2)
    doc.add_paragraph("Total estimated budget: $480,000")
    doc.add_paragraph("Engineering: $310,000 | Design: $60,000 | QA & Security: $80,000 | Contingency: $30,000")

    doc.save(HERE / "Project_Proposal_v3.docx")

def make_contract_docx():
    doc = Document()

    doc.add_heading("Service Agreement", level=1)
    doc.add_paragraph(
        "This Service Agreement (\"Agreement\") is entered into as of April 1, 2024, "
        "between AFSAM Technologies Inc. (\"Provider\") and Meridian Capital Group LLC "
        "(\"Client\")."
    )

    doc.add_heading("1. Scope of Services", level=2)
    doc.add_paragraph(
        "Provider agrees to supply the Client with access to the AFSAM financial analysis "
        "platform, including all standard modules, API access, and dedicated support "
        "as described in Schedule A."
    )

    doc.add_heading("2. Term", level=2)
    doc.add_paragraph(
        "This Agreement shall commence on April 1, 2024, and continue for a period of "
        "twelve (12) months, automatically renewing unless terminated with 30 days notice."
    )

    doc.add_heading("3. Fees and Payment", level=2)
    doc.add_paragraph(
        "Client agrees to pay $48,000 annually, billed quarterly at $12,000 per quarter, "
        "due within 30 days of invoice date. Late payments incur 1.5% monthly interest."
    )

    doc.add_heading("4. Confidentiality", level=2)
    doc.add_paragraph(
        "Both parties agree to maintain strict confidentiality of all proprietary "
        "information exchanged during the term of this Agreement and for five (5) years "
        "thereafter."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Signed:")
    doc.add_paragraph("_________________________          _________________________")
    doc.add_paragraph("AFSAM Technologies Inc.            Meridian Capital Group LLC")
    doc.add_paragraph("Sarah Kim, CTO                     Richard Holt, CEO")

    doc.save(HERE / "Service_Agreement_Meridian.docx")

make_proposal_docx()
make_contract_docx()

# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

files = sorted(p for p in HERE.iterdir() if p.name != "generate.py")
print("Generated files in test_victim/:")
for f in files:
    size_kb = f.stat().st_size / 1024
    print(f"  {f.name:<40} {size_kb:>6.1f} KB")
print(f"\nTotal: {len(files)} files ready for demo.")
print("Run this script again at any time to restore files after an attack demo.")
